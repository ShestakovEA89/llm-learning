from docx import Document

from generate_act_final import generate_act


def _all_docx_text(path):
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_generate_act_creates_docx_with_act_data(full_act, tmp_path):
    output_path = tmp_path / "test_act.docx"

    generate_act(act_id=full_act, output_path=str(output_path))

    assert output_path.exists()
    assert output_path.stat().st_size > 0

    text = _all_docx_text(str(output_path))
    assert "PYTEST-ACT-FULL" in text
    assert "Тестовые работы для генерации акта" in text
    assert "Тестовый Производитель Работ" in text
